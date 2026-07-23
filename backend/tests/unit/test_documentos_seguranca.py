from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
import uuid
import zipfile

import pytest
from fastapi import UploadFile

from app.core.config import settings
from app.modules.auth.dependencies import CurrentUser
from app.modules.documentos.exceptions import DocumentoInvalido, SanitizacaoFalhou
from app.modules.documentos import service as documentos_service
from app.modules.documentos.sanitizer import sanitizar_arquivo
from app.modules.documentos.validation import (
    detectar_formato,
    normalizar_nome_original,
    validar_extensao,
)


def test_nome_original_remove_caminho_e_controles() -> None:
    assert normalizar_nome_original("../../pasta\\teste\x00 clinico.pdf") == "teste clinico.pdf"
    assert normalizar_nome_original("laudo\u202egnp.exe.pdf") == "laudognp.exe.pdf"


@pytest.mark.parametrize("nome", [None, "", "...", "a" * 256 + ".pdf"])
def test_nome_original_invalido_falha_fechado(nome: str | None) -> None:
    with pytest.raises(DocumentoInvalido):
        normalizar_nome_original(nome)


def test_assinatura_real_nao_confia_na_extensao(tmp_path: Path) -> None:
    falso = tmp_path / "falso.pdf"
    falso.write_bytes(b"nao e um pdf")
    with pytest.raises(DocumentoInvalido):
        detectar_formato(falso)


def test_zip_so_e_aberto_dentro_do_worker(tmp_path: Path) -> None:
    falso = tmp_path / "falso.docx"
    falso.write_bytes(b"PK\x03\x04central-directory-maliciosa")
    assert detectar_formato(falso) == "docx"
    with pytest.raises(SanitizacaoFalhou):
        sanitizar_arquivo(falso, tmp_path / "saida.docx", "docx")


def test_extensao_precisa_corresponder_ao_conteudo() -> None:
    with pytest.raises(DocumentoInvalido):
        validar_extensao("imagem.pdf", "png")
    validar_extensao("foto.jpeg", "jpeg")


def test_docx_com_relacionamento_externo_e_recusado(tmp_path: Path) -> None:
    entrada = tmp_path / "externo.docx"
    with zipfile.ZipFile(entrada, "w", zipfile.ZIP_DEFLATED) as pacote:
        pacote.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
        )
        pacote.writestr("word/document.xml", "<document/>")
        pacote.writestr(
            "word/_rels/document.xml.rels",
            '<Relationships><Relationship TargetMode="External" Target="https://exemplo"/>'
            "</Relationships>",
        )
    saida = tmp_path / "saida.docx"
    with pytest.raises(SanitizacaoFalhou):
        sanitizar_arquivo(entrada, saida, "docx")
    assert not saida.exists()


def test_png_e_regravado_sem_metadados(tmp_path: Path) -> None:
    from PIL import Image, PngImagePlugin

    entrada = tmp_path / "entrada.png"
    metadata = PngImagePlugin.PngInfo()
    metadata.add_text("autor", "dado sensivel")
    Image.new("RGB", (8, 8), "red").save(entrada, pnginfo=metadata)
    saida = tmp_path / "saida.png"

    sanitizar_arquivo(entrada, saida, "png")

    with Image.open(saida) as imagem:
        assert imagem.format == "PNG"
        assert "autor" not in imagem.info
        assert imagem.size == (8, 8)


def test_pdf_simples_e_reconstruido(tmp_path: Path) -> None:
    import pikepdf

    entrada = tmp_path / "entrada.pdf"
    with pikepdf.new() as pdf:
        pdf.add_blank_page(page_size=(200, 200))
        pdf.docinfo["/Author"] = "dado sensivel"
        pdf.Root.OpenAction = pikepdf.Dictionary(
            S=pikepdf.Name("/JavaScript"), JS=pikepdf.String("app.alert('x')")
        )
        pdf.save(entrada)
    saida = tmp_path / "saida.pdf"

    sanitizar_arquivo(entrada, saida, "pdf")

    with pikepdf.open(saida) as pdf:
        assert len(pdf.pages) == 1
        assert "/Author" not in pdf.docinfo
        assert "/OpenAction" not in pdf.Root


def test_docx_e_reconstruido_com_texto_e_tabela(tmp_path: Path) -> None:
    from docx import Document

    entrada = tmp_path / "entrada.docx"
    documento = Document()
    documento.add_paragraph("Resultado do teste")
    tabela = documento.add_table(rows=1, cols=2)
    tabela.cell(0, 0).text = "Indice"
    tabela.cell(0, 1).text = "42"
    documento.core_properties.author = "dado sensivel"
    documento.save(entrada)
    saida = tmp_path / "saida.docx"

    sanitizar_arquivo(entrada, saida, "docx")

    limpo = Document(saida)
    assert "Resultado do teste" in [p.text for p in limpo.paragraphs]
    assert limpo.tables[0].cell(0, 1).text == "42"
    assert limpo.core_properties.author != "dado sensivel"


def test_docx_recusa_elemento_que_seria_perdido(tmp_path: Path) -> None:
    from docx import Document

    entrada = tmp_path / "comentarios.docx"
    base = tmp_path / "base.docx"
    Document().save(base)
    with zipfile.ZipFile(base) as pacote:
        partes = {info.filename: pacote.read(info) for info in pacote.infolist()}
    with zipfile.ZipFile(entrada, "w", zipfile.ZIP_DEFLATED) as pacote:
        for nome, conteudo in partes.items():
            pacote.writestr(nome, conteudo)
        pacote.writestr("word/comments.xml", "<w:comments/>")

    with pytest.raises(SanitizacaoFalhou):
        sanitizar_arquivo(entrada, tmp_path / "saida.docx", "docx")


def test_docx_nao_expoe_imagem_orfa_do_pacote(tmp_path: Path) -> None:
    from docx import Document
    from PIL import Image

    base = tmp_path / "base.docx"
    documento = Document()
    documento.add_paragraph("Conteudo visivel")
    documento.save(base)
    imagem = BytesIO()
    Image.new("RGB", (4, 4), "blue").save(imagem, "PNG")

    entrada = tmp_path / "orfa.docx"
    with zipfile.ZipFile(base) as pacote:
        partes = {info.filename: pacote.read(info) for info in pacote.infolist()}
    tipos = partes["[Content_Types].xml"].replace(
        b"</Types>",
        b'<Default Extension="png" ContentType="image/png"/></Types>',
    )
    partes["[Content_Types].xml"] = tipos
    partes["word/media/orfa.png"] = imagem.getvalue()
    with zipfile.ZipFile(entrada, "w", zipfile.ZIP_DEFLATED) as pacote:
        for nome, conteudo in partes.items():
            pacote.writestr(nome, conteudo)

    saida = tmp_path / "saida.docx"
    sanitizar_arquivo(entrada, saida, "docx")
    limpo = Document(saida)
    assert [p.text for p in limpo.paragraphs] == ["Conteudo visivel"]


def test_fluxo_de_upload_grava_somente_resultado_sanitizado(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from PIL import Image, PngImagePlugin

    paciente_id = uuid.uuid4()
    tenant_id = uuid.uuid4()
    user = CurrentUser(id=uuid.uuid4(), tenant_id=tenant_id, papel="psicologa")
    paciente = SimpleNamespace(id=paciente_id)

    class Resultado:
        def scalar_one_or_none(self):
            return paciente

        def scalar_one(self):
            return 0

    class BancoFalso:
        adicionados: list[object] = []

        def get(self, _modelo, _id):
            return paciente

        def execute(self, _stmt, _params=None):
            return Resultado()

        def add(self, item):
            self.adicionados.append(item)

        def flush(self):
            return None

    raiz = tmp_path / "privado"
    monkeypatch.setattr(settings, "documentos_dir", str(raiz))
    monkeypatch.setattr(documentos_service.event, "listen", lambda *_args, **_kwargs: None)
    monkeypatch.setattr(documentos_service.audit_service, "registrar_evento", lambda *_a, **_k: None)

    bruto = BytesIO()
    metadata = PngImagePlugin.PngInfo()
    metadata.add_text("autor", "nao deve persistir")
    Image.new("RGB", (6, 6), "green").save(bruto, "PNG", pnginfo=metadata)
    bruto.seek(0)
    upload = UploadFile(filename="resultado.png", file=bruto)

    documento = documentos_service.enviar(BancoFalso(), user, paciente_id, upload)

    assert documento is not None
    armazenado = documentos_service._resolver_chave(documento.chave_armazenamento)
    assert armazenado.is_file()
    with Image.open(armazenado) as imagem:
        assert "autor" not in imagem.info
    assert documento.sha256 == documentos_service._sha256(armazenado)
    assert not any(caminho.name == "entrada" for caminho in raiz.rglob("*"))
