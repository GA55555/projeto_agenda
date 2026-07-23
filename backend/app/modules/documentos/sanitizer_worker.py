"""Worker SEM acesso ao app/BD: reconstrui PDF, DOCX, JPEG ou PNG.

Executado com `python -I`, ambiente sem credenciais, limites de SO e timeout no
processo pai. Nao abre sockets e nunca usa shell.
"""
from __future__ import annotations

import os
from pathlib import Path, PurePosixPath
import sys
import tempfile
import warnings
import zipfile

MAX_PIXELS = 20_000_000
MAX_DOCX_ENTRIES = 2_000
MAX_DOCX_UNCOMPRESSED = 100 * 1024 * 1024
MAX_DOCX_RATIO = 100
MAX_PDF_PAGES = 500


def _limitar_recursos() -> None:
    if os.name != "posix":
        return
    import resource

    memoria = 384 * 1024 * 1024
    resource.setrlimit(resource.RLIMIT_AS, (memoria, memoria))
    resource.setrlimit(resource.RLIMIT_CPU, (25, 30))
    resource.setrlimit(resource.RLIMIT_FSIZE, (64 * 1024 * 1024, 64 * 1024 * 1024))
    resource.setrlimit(resource.RLIMIT_NOFILE, (64, 64))


def _sanitizar_imagem(entrada: Path, saida: Path, formato: str) -> None:
    from PIL import Image

    Image.MAX_IMAGE_PIXELS = MAX_PIXELS
    warnings.simplefilter("error", Image.DecompressionBombWarning)
    permitido = ["JPEG"] if formato == "jpeg" else ["PNG"]
    with Image.open(entrada, formats=permitido) as origem:
        origem.load()
        if getattr(origem, "n_frames", 1) != 1:
            raise ValueError("Imagem com multiplos quadros nao e aceita")
        if origem.width <= 0 or origem.height <= 0:
            raise ValueError("Dimensoes invalidas")
        if formato == "jpeg":
            pixels = origem.convert("RGB")
            limpa = Image.new("RGB", pixels.size, "white")
            limpa.paste(pixels)
            limpa.save(saida, "JPEG", quality=92, optimize=False, exif=b"")
        else:
            modo = "RGBA" if "A" in origem.getbands() else "RGB"
            pixels = origem.convert(modo)
            limpa = Image.new(modo, pixels.size)
            limpa.paste(pixels)
            limpa.save(saida, "PNG", compress_level=6)


def _sanitizar_pdf(entrada: Path, saida: Path) -> None:
    import pikepdf
    from pikepdf import sanitize

    with pikepdf.open(entrada) as pdf:
        if pdf.is_encrypted:
            raise ValueError("PDF protegido por senha nao e aceito")
        if len(pdf.pages) > MAX_PDF_PAGES:
            raise ValueError(f"PDF excede {MAX_PDF_PAGES} paginas")
        sanitize.remove_javascript(pdf)
        sanitize.remove_attachments(pdf)
        sanitize.remove_external_access(pdf)
        sanitize.remove_thumbnails(pdf)
        sanitize.remove_search_index(pdf)
        sanitize.remove_multimedia(pdf)
        sanitize.remove_web_capture(pdf)
        sanitize.remove_private_app_data(pdf)
        sanitize.remove_collection(pdf)
        for chave in list(pdf.docinfo.keys()):
            del pdf.docinfo[chave]
        if "/Metadata" in pdf.Root:
            del pdf.Root["/Metadata"]
        pdf.save(
            saida,
            object_stream_mode=pikepdf.ObjectStreamMode.generate,
            recompress_flate=True,
            deterministic_id=True,
        )


def _validar_docx(entrada: Path) -> None:
    with zipfile.ZipFile(entrada) as pacote:
        infos = pacote.infolist()
        if not infos or len(infos) > MAX_DOCX_ENTRIES:
            raise ValueError("Quantidade invalida de partes no DOCX")
        total = 0
        vistos: set[str] = set()
        for info in infos:
            nome = info.filename
            caminho = PurePosixPath(nome)
            chave_nome = nome.casefold()
            if (
                not nome
                or "\\" in nome
                or caminho.is_absolute()
                or ".." in caminho.parts
                or (info.flag_bits & 0x1)
                or chave_nome in vistos
            ):
                raise ValueError("DOCX contem caminho, duplicata ou criptografia invalida")
            vistos.add(chave_nome)
            total += info.file_size
            if total > MAX_DOCX_UNCOMPRESSED:
                raise ValueError("DOCX expande alem do limite permitido")
            if info.file_size and info.compress_size == 0:
                raise ValueError("DOCX contem entrada de compressao suspeita")
            if info.compress_size and info.file_size / info.compress_size > MAX_DOCX_RATIO:
                raise ValueError("DOCX contem possivel ZIP bomb")
            baixo = nome.lower()
            if any(token in baixo for token in ("vbaproject", "/activex/", "/embeddings/")):
                raise ValueError("DOCX contem macro ou objeto incorporado")
        tipos = pacote.read("[Content_Types].xml").lower()
        if b"macroenabled" in tipos or b"vba" in tipos:
            raise ValueError("DOCX habilitado para macros")
        for info in infos:
            if not info.filename.lower().endswith(".rels"):
                continue
            xml = pacote.read(info)
            baixo = xml.lower()
            if b'targetmode="external"' in baixo or b"targetmode='external'" in baixo:
                raise ValueError("DOCX contem relacionamento externo")
            if b"oleobject" in baixo or b"attachedtemplate" in baixo:
                raise ValueError("DOCX contem objeto ativo")

        # O reconstrutor preserva paragrafos, tabelas, cabecalhos, rodapes e
        # imagens. Elementos que ele nao representa sao recusados para nao perder
        # conteudo clinico silenciosamente.
        nomes_baixos = {info.filename.lower() for info in infos}
        nao_suportados = {
            "word/comments.xml",
            "word/footnotes.xml",
            "word/endnotes.xml",
        }
        if nomes_baixos & nao_suportados:
            raise ValueError("DOCX contem comentarios, notas ou elementos nao suportados")
        tags_nao_suportadas = (
            b"<w:altchunk",
            b"<w:object",
            b"<w:pict",
            b"<w:txbxcontent",
            b"<w:sdt",
            b"<w:ins",
            b"<w:del",
            b"<mc:alternatecontent",
            b"<m:omath",
        )
        partes_textuais = [
            info
            for info in infos
            if info.filename == "word/document.xml"
            or (
                info.filename.lower().endswith(".xml")
                and info.filename.lower().startswith(("word/header", "word/footer"))
            )
        ]
        for parte in partes_textuais:
            xml = pacote.read(parte).lower()
            if any(tag in xml for tag in tags_nao_suportadas):
                raise ValueError("DOCX contem estrutura que nao pode ser reconstruida sem perdas")


def _copiar_paragrafo(origem, destino) -> None:
    novo = destino.add_paragraph()
    for run in origem.runs:
        trecho = novo.add_run(run.text)
        trecho.bold = run.bold
        trecho.italic = run.italic
        trecho.underline = run.underline
    if novo.text != origem.text:
        # `Paragraph.runs` pode omitir texto dentro de hyperlinks. Nesse caso,
        # prioriza integridade do conteudo sobre a formatacao parcial.
        for filho in list(novo._element):
            novo._element.remove(filho)
        novo.add_run(origem.text)


def _copiar_tabela(origem, destino) -> None:
    if not origem.rows:
        return
    colunas = max(len(linha.cells) for linha in origem.rows)
    tabela = destino.add_table(rows=0, cols=colunas)
    tabela.style = "Table Grid"
    for linha_origem in origem.rows:
        celulas = tabela.add_row().cells
        for indice, celula in enumerate(linha_origem.cells):
            celulas[indice].text = celula.text


def _sanitizar_docx(entrada: Path, saida: Path) -> None:
    from docx import Document
    from docx.parts.image import ImagePart
    from docx.shared import Inches
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    _validar_docx(entrada)
    origem = Document(entrada)
    destino = Document()
    # Remove o paragrafo vazio criado pelo template padrao.
    if destino.paragraphs:
        elemento = destino.paragraphs[0]._element
        elemento.getparent().remove(elemento)
    for bloco in origem.iter_inner_content():
        if isinstance(bloco, Paragraph):
            _copiar_paragrafo(bloco, destino)
        elif isinstance(bloco, Table):
            _copiar_tabela(bloco, destino)

    extras: list[tuple[str, list[str]]] = []
    for secao in origem.sections:
        cab = [p.text for p in secao.header.paragraphs if p.text.strip()]
        rod = [p.text for p in secao.footer.paragraphs if p.text.strip()]
        if cab:
            extras.append(("Cabecalho", cab))
        if rod:
            extras.append(("Rodape", rod))
    for titulo, textos in extras:
        destino.add_heading(titulo, level=2)
        for texto in textos:
            destino.add_paragraph(texto)

    # Somente ImageParts alcancaveis pelas relacoes do documento; midia orfa no
    # ZIP nao ganha visibilidade no arquivo reconstruido.
    imagens = [parte for parte in origem.part.package.parts if isinstance(parte, ImagePart)]
    if imagens:
        destino.add_heading("Imagens do documento", level=2)
        with tempfile.TemporaryDirectory() as temporario:
            pasta = Path(temporario)
            for indice, parte in enumerate(imagens):
                bruto = pasta / f"origem-{indice}"
                bruto.write_bytes(parte.blob)
                with bruto.open("rb") as arquivo:
                    inicio = arquivo.read(8)
                if inicio.startswith(b"\x89PNG\r\n\x1a\n"):
                    formato, sufixo = "png", ".png"
                elif inicio.startswith(b"\xff\xd8\xff"):
                    formato, sufixo = "jpeg", ".jpg"
                else:
                    raise ValueError("DOCX contem imagem fora da allowlist JPEG/PNG")
                limpa = pasta / f"limpa-{indice}{sufixo}"
                _sanitizar_imagem(bruto, limpa, formato)
                destino.add_picture(str(limpa), width=Inches(6.5))
    destino.save(saida)


def main() -> int:
    _limitar_recursos()
    if len(sys.argv) != 4:
        return 2
    entrada, saida, formato = Path(sys.argv[1]), Path(sys.argv[2]), sys.argv[3]
    if formato in {"jpeg", "png"}:
        _sanitizar_imagem(entrada, saida, formato)
    elif formato == "pdf":
        _sanitizar_pdf(entrada, saida)
    elif formato == "docx":
        _sanitizar_docx(entrada, saida)
    else:
        raise ValueError("Formato nao permitido")
    os.chmod(saida, 0o600)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:  # mensagem curta; nunca inclui conteudo do arquivo
        print(f"{type(exc).__name__}: {exc}", file=sys.stderr)
        raise SystemExit(1) from None
